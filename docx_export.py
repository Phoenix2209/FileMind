# docx_export.py
# Converts a SummaryReport into a clean, readable Word document (.docx)
# Uses python-docx — works fully offline, no internet required.

from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from schema import SummaryReport


# ── Colours ────────────────────────────────────────────────────────────────
ACCENT_RGB = RGBColor(0x4F, 0x8E, 0xF7)
MUTED_RGB  = RGBColor(0x64, 0x74, 0x8B)
GREEN_RGB  = RGBColor(0x3E, 0xCF, 0x8E)
RED_RGB    = RGBColor(0xC0, 0x39, 0x2B)


def export_to_docx(report: SummaryReport, output_path: Path) -> Path:
    """
    Builds a Word document from a SummaryReport and saves it.
    Returns the path to the saved file.
    """
    doc = Document()
    _set_default_font(doc)

    _add_title_page(doc, report)
    _add_global_summary(doc, report)
    _add_themes_and_insights(doc, report)
    _add_per_file_table(doc, report)
    if report.failed:
        _add_failed_files(doc, report)

    doc.save(str(output_path))
    return output_path


# ── Sections ───────────────────────────────────────────────────────────────

def _set_default_font(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _add_title_page(doc: Document, report: SummaryReport):
    title = doc.add_heading("FileMind — Summary Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("AI-powered offline file summarizer")
    run.italic = True
    run.font.color.rgb = MUTED_RGB
    run.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # Prompt used
    p = doc.add_paragraph()
    p.add_run("Task: ").bold = True
    p.add_run(report.prompt)

    # Stats line
    stats = doc.add_paragraph()
    stats.add_run("Files processed: ").bold = True
    stats.add_run(f"{report.processed} / {report.total_files}")
    if report.failed:
        stats.add_run(f"   (Failed: {report.failed})").font.color.rgb = RED_RGB

    model_p = doc.add_paragraph()
    model_p.add_run("Model used: ").bold = True
    model_p.add_run(report.model_used)

    doc.add_paragraph()  # spacer
    _add_divider(doc)


def _add_global_summary(doc: Document, report: SummaryReport):
    doc.add_heading("Global Summary", level=1)
    gs = report.global_summary.global_summary
    if gs:
        doc.add_paragraph(gs)
    else:
        p = doc.add_paragraph()
        p.add_run("No global summary was generated.").italic = True
    doc.add_paragraph()


def _add_themes_and_insights(doc: Document, report: SummaryReport):
    themes   = report.global_summary.common_themes
    insights = report.global_summary.top_insights

    if themes:
        doc.add_heading("Common Themes", level=2)
        for t in themes:
            doc.add_paragraph(t, style="List Bullet")
        doc.add_paragraph()

    if insights:
        doc.add_heading("Top Insights", level=2)
        for i in insights:
            doc.add_paragraph(i, style="List Bullet")
        doc.add_paragraph()


def _add_per_file_table(doc: Document, report: SummaryReport):
    doc.add_heading("Per-File Breakdown", level=1)

    successful_files = [f for f in report.files if f.success]
    if not successful_files:
        doc.add_paragraph("No files were successfully processed.")
        return

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0].cells
    headers = ["File", "Data Summary", "Key Points", "Entities"]
    for cell, text in zip(hdr, headers):
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True

    # Column widths
    widths = [Inches(1.3), Inches(2.2), Inches(2.5), Inches(1.5)]
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w

    for fs in successful_files:
        row = table.add_row().cells

        row[0].text = ""
        p0 = row[0].paragraphs[0]
        p0.add_run(fs.filename).bold = True
        p0.add_run(f"\n{fs.file_type}  ·  {fs.file_size_kb} KB").italic = True

        row[1].text = fs.data_summary or "—"

        row[2].text = ""
        p2 = row[2].paragraphs[0]
        if fs.key_points:
            for idx, kp in enumerate(fs.key_points):
                if idx == 0:
                    p2.add_run(f"• {kp}")
                else:
                    row[2].add_paragraph(f"• {kp}")
        else:
            p2.add_run("—")

        row[3].text = fs.entities or "—"

    doc.add_paragraph()


def _add_failed_files(doc: Document, report: SummaryReport):
    doc.add_heading("Failed Files", level=1)
    for fs in report.files:
        if not fs.success:
            p = doc.add_paragraph()
            run = p.add_run(f"✗ {fs.filename}: ")
            run.bold = True
            run.font.color.rgb = RED_RGB
            p.add_run(fs.error or "Unknown error")


def _add_divider(doc: Document):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)